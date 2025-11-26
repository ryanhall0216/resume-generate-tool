import streamlit as st
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
from reportlab.lib import colors
from openai import OpenAI
import json
import os
from dotenv import load_dotenv
import time
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Load env
load_dotenv()
api_key = os.getenv("OPENAI_API_KEY")
client = OpenAI(api_key=api_key)

# ----------------------- helpers -----------------------

DEFAULT_PROMPT = """

Rules:
- Tailor resume. All required but also bonus skills are that I have hands-on exp.
- Add even if not in the current resume.
- Give 15-20 bullets per position.
- Each bullet must be long enough and should be meaningful by containing 20-40 words.
- When add technologies per position, at least 20-40 technologies should be added.
- Add all required and bonus skills or reflect behavior in the latest position.
- Add some core skills or behaviors in the other positions as well.
- Add education from the resume template.
- Remove unnecessary skills.
- Always fill the Header fields from the resume

"""

def build_header_from_json(tailored: dict):
    """
    Build (name_line, contact_line) from the model's JSON.
    Expected:
    tailored["Header"] = {
        "Name": str,
        "Email": str (optional),
        "Phone": str (optional),
        "Address": str (optional),
        "Links": [str, ...] (optional),
        "Other": [str, ...] (optional)  # any other bits like GitHub/LinkedIn if not in Links
    }
    """
    header = tailored.get("Header", {}) or {}
    name = (header.get("Name") or "").strip()

    # Collect contact parts in a consistent order
    parts = []
    if header.get("Address"):
        parts.append(header["Address"].strip())
    if header.get("Phone"):
        parts.append(header["Phone"].strip())
    if header.get("Email"):
        parts.append(header["Email"].strip())

    # Links/Other arrays
    links = header.get("Links") or []
    other = header.get("Other") or []
    for item in links:
        s = str(item).strip()
        if s:
            parts.append(s)
    for item in other:
        s = str(item).strip()
        if s:
            parts.append(s)

    contact_line = " • ".join([p for p in parts if p])
    return name, contact_line

def clean_bullet(text: str) -> str:
    return text.lstrip("•*-·–— ").strip()

def get_alignment(value: str):
    mapping_docx = {
        "Left": WD_ALIGN_PARAGRAPH.LEFT,
        "Center": WD_ALIGN_PARAGRAPH.CENTER,
        "Right": WD_ALIGN_PARAGRAPH.RIGHT,
        "Justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    mapping_pdf = {
        "Left": TA_LEFT,
        "Center": TA_CENTER,
        "Right": TA_RIGHT,
        "Justify": TA_JUSTIFY,
    }
    return mapping_docx.get(value, WD_ALIGN_PARAGRAPH.LEFT), mapping_pdf.get(value, TA_LEFT)

def tailor_resume(resume_text, job_description, custom_prompt=None):
    base_prompt = """
    You are a resume expert. Tailor the given resume to the job description.

    Keep the structure JSON with these keys only:
    - Summary
    - Skills (dict of categories: items[])
    - Experience (list of jobs, each with Company, Title, Dates, Responsibilities[], Technologies is optional)
    - Education (must always include: Dates, Degree, Institution, GPA, Institution = university/school name, Location is optional)

    You are a resume expert. Tailor the given resume to the job description.

    Return a STRICT JSON object with these keys ONLY:
    - Header (object):
        - Name (string, required)
        - Email (string, if known)
        - Phone (string, if known)
        - Address (string, if known)
        - Links (array of strings, optional; e.g., LinkedIn, GitHub, portfolio)
        - Other (array of strings, optional; e.g., extra contact info)
    - Summary (string)
    - Skills (object where each key is a category and value is an array of strings)
    - Experience (array of objects, each with: Company, Title, Dates, Responsibilities[]; Technologies is optional array)
    - Education (object where the keys are Dates, Degree, Institution, GPA; Location is optional)

    Constraints:
    - Be truthful to resume and job description.
    - If any contact field is missing, omit that field rather than inventing.
    """

    if custom_prompt and custom_prompt.strip():
        prompt = base_prompt + "\n\n" + custom_prompt.strip()
    else:
        prompt = base_prompt

    response = client.chat.completions.create(
        model="gpt-4.1-mini",
        messages=[
            {"role": "system", "content": "You are a helpful assistant that ONLY returns JSON when asked."},
            {"role": "user", "content": f"(Return JSON only)\n{prompt}\n\nResume:\n{resume_text}\n\nJob Description:\n{job_description}"}
        ],
        response_format={"type": "json_object"}
    )
    return json.loads(response.choices[0].message.content)

def _pt_to_px(pt):
    # rough screen px conversion for preview
    return int(round(pt * 1.333))

def _align_to_css(align_key: str) -> str:
    return {
        "Left": "left",
        "Center": "center",
        "Right": "right",
        "Justify": "justify",
    }.get(align_key, "left")

def render_preview_html(tailored: dict, styles_config: dict) -> str:
    # read header
    name, contact = ("", "")
    if tailored:
        header = (tailored.get("Header") or {})
        name = (header.get("Name") or "").strip()
        parts = []
        if header.get("Address"): parts.append(header["Address"])
        if header.get("Phone"):   parts.append(header["Phone"])
        if header.get("Email"):   parts.append(header["Email"])
        for v in header.get("Links") or []: parts.append(str(v))
        for v in header.get("Other") or []: parts.append(str(v))
        contact = " • ".join([p for p in parts if p])

    # sizes/alignment
    title_css    = f"font-size:{_pt_to_px(styles_config['title_size'])}px;font-weight:{'700' if styles_config['title_bold'] else '400'};font-style:{'italic' if styles_config['title_italic'] else 'normal'};text-align:{_align_to_css(styles_config['title_align'])};margin:0;"
    contact_css  = f"font-size:{_pt_to_px(styles_config['body_size'])}px;font-weight:{'700' if styles_config['contact_bold'] else '400'};font-style:{'italic' if styles_config['contact_italic'] else 'normal'};text-align:{_align_to_css(styles_config['contact_align'])};margin:2px 0 12px 0;color:#444;"
    heading_css  = f"font-size:{_pt_to_px(styles_config['heading_size'])}px;font-weight:{'700' if styles_config['heading_bold'] else '400'};font-style:{'italic' if styles_config['heading_italic'] else 'normal'};text-align:{_align_to_css(styles_config['heading_align'])};margin:16px 0 6px 0;border-bottom:1px solid #eee;padding-bottom:2px;"
    sub_css      = f"font-size:{_pt_to_px(styles_config['subheading_size'])}px;font-weight:{'700' if styles_config['subheading_bold'] else '400'};font-style:{'italic' if styles_config['subheading_italic'] else 'normal'};text-align:{_align_to_css(styles_config['subheading_align'])};margin:10px 0 4px 0;color:#333;"
    body_css     = f"font-size:{_pt_to_px(styles_config['body_size'])}px;font-weight:{'700' if styles_config['body_bold'] else '400'};font-style:{'italic' if styles_config['body_italic'] else 'normal'};text-align:{_align_to_css(styles_config['body_align'])};margin:2px 0;color:#222;"

    # sections
    summary = (tailored or {}).get("Summary") or ""
    skills  = (tailored or {}).get("Skills") or {}
    exp     = (tailored or {}).get("Experience") or []
    edu     = (tailored or {}).get("Education")

    # education: your prompt returns an object — render that. If later you make it a list, handle both.
    edu_html = ""
    if isinstance(edu, dict):
        lines = []
        if edu.get("Dates"): lines.append(edu["Dates"])
        degree_bits = []
        if edu.get("Degree"): degree_bits.append(edu["Degree"])
        for key in ("Institution","University","School"):
            if edu.get(key):
                degree_bits.append(edu[key]); break
        if edu.get("Location"): degree_bits.append(edu["Location"])
        if degree_bits: lines.append(" | ".join(degree_bits))
        if edu.get("GPA"): lines.append(f"GPA: {edu['GPA']}")
        edu_html = "".join([f'<div style="{body_css}">{line}</div>' for line in lines])

    # skills list
    skills_html = "".join(
        [f'<div style="{body_css}">• {cat}: {", ".join(items)}</div>' for cat, items in skills.items()]
    )

    # experience (limit bullets for speed)
    exp_blocks = []
    for job in exp[:4]:  # show first 4 roles in preview
        header_line = f"{job.get('Title','')} | {job.get('Company','')} | {job.get('Dates','')}"
        bullets = job.get("Responsibilities") or []
        bullets_html = "".join([f'<li><span style="{body_css}">{str(b).lstrip("•*-·–— ").strip()}</span></li>' for b in bullets[:8]])  # first 8 bullets
        techs_html = ""
        if isinstance(job.get("Technologies"), list) and job["Technologies"]:
            techs_html = f'<div style="{body_css}"><em>Technologies:</em> {", ".join(job["Technologies"][:40])}</div>'
        exp_blocks.append(
            f'<div><div style="{sub_css}">{header_line}</div><ul style="margin:0 0 6px 20px;padding:0;">{bullets_html}</ul>{techs_html}</div>'
        )
    exp_html = "".join(exp_blocks)

    # final HTML
    html = f"""
    <div style="font-family: Arial, sans-serif; line-height:1.35;">
      <div style="{title_css}">{name}</div>
      {f'<div style="{contact_css}">{contact}</div>' if contact else ''}
      <div style="{heading_css}">Summary</div>
      <div style="{body_css}">{summary}</div>

      <div style="{heading_css}">Skills</div>
      {skills_html}

      <div style="{heading_css}">Experience</div>
      {exp_html}

      <div style="{heading_css}">Education</div>
      {edu_html}
    </div>
    """
    return html

# ----------------------- DOCX -----------------------

def create_docx(template_text, tailored, output_path, include_tech, styles_config):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    name_line, contact_text = build_header_from_json(tailored)

    # Title
    name_p = doc.add_paragraph(name_line)
    run = name_p.runs[0]
    run.font.size = Pt(styles_config["title_size"])
    run.font.name = 'Speak Pro (Headings)'
    run.bold = styles_config["title_bold"]
    run.italic = styles_config["title_italic"]
    name_p.alignment, _ = get_alignment(styles_config["title_align"])
    name_p.paragraph_format.space_after = Pt(0)

    # Contact
    if contact_text:
        c = doc.add_paragraph(contact_text)
        run = c.runs[0]
        run.font.size = Pt(styles_config["body_size"])
        run.font.name = 'Arial'
        run.bold = styles_config["contact_bold"]
        run.italic = styles_config["contact_italic"]
        c.alignment, _ = get_alignment(styles_config["contact_align"])

    # Section helpers
    def add_heading(text):
        p = doc.add_paragraph(text)
        r = p.runs[0]
        r.bold = styles_config["heading_bold"]
        r.font.name = 'Speak Pro (Headings)'
        r.italic = styles_config["heading_italic"]
        r.font.size = Pt(styles_config["heading_size"])
        p.alignment, _ = get_alignment(styles_config["heading_align"])

    def add_body(text, style=None):
        p = doc.add_paragraph(text, style=style)
        r = p.runs[0]
        r.bold = styles_config["body_bold"]
        r.font.name = 'Arial'
        r.italic = styles_config["body_italic"]
        r.font.size = Pt(styles_config["body_size"])
        p.alignment, _ = get_alignment(styles_config["body_align"])
        return p

    # Summary
    add_heading("Summary")
    add_body(tailored["Summary"])

    # Skills
    add_heading("Skills")
    for category, items in tailored["Skills"].items():
        add_body(f"{category}: {', '.join(items)}", style="List Bullet")

    # Experience
    add_heading("Experience")
    for job in tailored["Experience"]:
        heading = doc.add_paragraph(f"{job['Title']} | {job['Company']} | {job['Dates']}")
        run = heading.runs[0]
        run.bold = styles_config["subheading_bold"]
        run.italic = styles_config["subheading_italic"]
        run.font.size = Pt(styles_config["subheading_size"])
        heading.alignment, _ = get_alignment(styles_config["subheading_align"])
        for resp in job["Responsibilities"]:
            p = add_body(clean_bullet(resp), style="List Bullet")
            p.paragraph_format.space_after = Pt(0)
        if include_tech and "Technologies" in job:
            techs = ", ".join(job["Technologies"]) if isinstance(job["Technologies"], list) else str(job["Technologies"])
            add_body(f"Technologies: {techs}")

    # Education
    add_heading("Education")
    edu = tailored["Education"]

    if isinstance(edu, dict):
        if "Dates" in edu:
            add_body(edu["Dates"])
        degree_line = []
        if "Degree" in edu:
            degree_line.append(edu["Degree"])
        for key in ("Institution", "University", "School"):
            if key in edu:
                degree_line.append(edu[key])
                break
        if "Location" in edu:
            degree_line.append(edu["Location"])
        if degree_line:
            add_body(" | ".join(degree_line))
        if "GPA" in edu:
            add_body(f"GPA: {edu['GPA']}")

    doc.save(output_path)
    return output_path

# ----------------------- PDF -----------------------

def create_pdf(template_text, tailored, output_path, include_tech, styles_config):
    def make_style(name, size, bold, italic, align_key):
        _, pdf_align = get_alignment(align_key)
        return ParagraphStyle(
            name,
            fontName="Helvetica-Bold" if bold else "Helvetica",
            fontSize=size,
            leading=size + 2,
            textColor=colors.black,
            alignment=pdf_align,
        )

    pdf_styles = {
        "Title": make_style("Title", styles_config["title_size"], styles_config["title_bold"], styles_config["title_italic"], styles_config["title_align"]),
        "Contact": make_style("Contact", styles_config["body_size"], styles_config["contact_bold"], styles_config["contact_italic"], styles_config["contact_align"]),
        "Heading": make_style("Heading", styles_config["heading_size"], styles_config["heading_bold"], styles_config["heading_italic"], styles_config["heading_align"]),
        "Subheading": make_style("Subheading", styles_config["subheading_size"], styles_config["subheading_bold"], styles_config["subheading_italic"], styles_config["subheading_align"]),
        "Body": make_style("Body", styles_config["body_size"], styles_config["body_bold"], styles_config["body_italic"], styles_config["body_align"]),
    }

    doc = SimpleDocTemplate(output_path)
    elements = []

    # Header
    name_line, contact_text = build_header_from_json(tailored)
    elements.append(Paragraph(name_line, pdf_styles["Title"]))
    if contact_text:
        elements.append(Paragraph(contact_text, pdf_styles["Contact"]))
    elements.append(Spacer(1, 12))

    # Summary
    elements.append(Paragraph("Summary", pdf_styles["Heading"]))
    elements.append(Paragraph(tailored["Summary"], pdf_styles["Body"]))
    elements.append(Spacer(1, 12))

    # Skills
    elements.append(Paragraph("Skills", pdf_styles["Heading"]))
    skill_items = [ListItem(Paragraph(f"{cat}: {', '.join(items)}", pdf_styles["Body"])) for cat, items in tailored["Skills"].items()]
    elements.append(ListFlowable(skill_items, bulletType="bullet"))
    elements.append(Spacer(1, 12))

    # Experience
    elements.append(Paragraph("Experience", pdf_styles["Heading"]))
    for job in tailored["Experience"]:
        elements.append(Paragraph(f"{job['Title']} | {job['Company']} | {job['Dates']}", pdf_styles["Subheading"]))
        bullets = [ListItem(Paragraph(clean_bullet(resp), pdf_styles["Body"])) for resp in job["Responsibilities"]]
        elements.append(ListFlowable(bullets, bulletType="bullet"))
        if include_tech and "Technologies" in job:
            techs = ", ".join(job["Technologies"]) if isinstance(job["Technologies"], list) else str(job["Technologies"])
            elements.append(Paragraph(f"Technologies: {techs}", pdf_styles["Body"]))
        elements.append(Spacer(1, 8))
    elements.append(Spacer(1, 12))

    # Education
    elements.append(Paragraph("Education", pdf_styles["Heading"]))
    edu = tailored["Education"]
    if isinstance(edu, dict):
        if "Dates" in edu:
            elements.append(Paragraph(edu["Dates"], pdf_styles["Body"]))
        degree_line = []
        if "Degree" in edu:
            degree_line.append(edu["Degree"])
        for key in ("Institution", "University", "School"):
            if key in edu:
                degree_line.append(edu[key])
                break
        if "Location" in edu:
            degree_line.append(edu["Location"])
        if degree_line:
            elements.append(Paragraph(" | ".join(degree_line), pdf_styles["Body"]))
        if "GPA" in edu:
            elements.append(Paragraph(f"GPA: {edu['GPA']}", pdf_styles["Body"]))

    doc.build(elements)
    return output_path

# ----------------------- Generate function -----------------------

def generate_resume():
    if st.session_state["loading"]:
        return
    st.session_state["loading"] = True
    st.session_state["stop_requested"] = False

    try:
        # ✅ spinner shows only during work and vanishes automatically on completion
        with st.spinner("Generating your tailored resume..."):

            if st.session_state["stop_requested"]:
                return

            tailored = tailor_resume(resume_template_text, job_description, custom_prompt)
            st.session_state["tailored_json"] = tailored

            if st.session_state["stop_requested"]:
                return

            os.makedirs("outputs", exist_ok=True)
            user_name, _ = build_header_from_json(tailored)
            user_name = user_name.replace(" ", "_") if user_name else "Software_Engineer"
            ts = int(time.time())

            if output_format == "DOCX":
                final_file = create_docx(resume_template_text, tailored, f"outputs/{user_name}_{ts}.docx", include_tech, styles_config)
            else:
                final_file = create_pdf(resume_template_text, tailored, f"outputs/{user_name}_{ts}.pdf", include_tech, styles_config)

            # Success + download button (this stays after completion; spinner/stop won’t)
            st.success(f"✅ Resume tailored successfully! ({output_format})")
            with open(final_file, "rb") as f:
                st.download_button(
                    f"Download Tailored Resume ({output_format})",
                    f,
                    file_name=f"{user_name}.{output_format.lower()}",
                )

    except Exception as e:
        if not st.session_state["stop_requested"]:
            st.error(f"❌ Something went wrong: {e}")
    finally:
        # ✅ turning this off hides the stop button + any loading UI in the main layout
        st.session_state["loading"] = False

# ----------------------- app -----------------------

st.set_page_config(layout="wide")

st.markdown("""
<style>
:root {
    --background-color: #ffffff;
    --secondary-background-color: #f7f7f9;
    --text-color: #000000;
}

[data-testid="stAppViewContainer"] {
    background-color: var(--background-color);
    color: var(--text-color);
}
[data-testid="stSidebar"] {
    background-color: var(--secondary-background-color);
}
</style>
""", unsafe_allow_html=True)

st.title("📄 Resume Tailor Tool")
st.sidebar.header("Settings")

if "loading" not in st.session_state:
    st.session_state["loading"] = False
if "stop_requested" not in st.session_state:
    st.session_state["stop_requested"] = False

disabled_state = st.session_state["loading"]

# Style settings
st.sidebar.subheader("🎨 Style Settings")
def style_section(label, defaults):
    st.sidebar.markdown(f"**{label}**")
    return {
        f"{label.lower()}_bold": st.sidebar.checkbox(f"{label} Bold", value=defaults["bold"], disabled=disabled_state),
        f"{label.lower()}_italic": st.sidebar.checkbox(f"{label} Italic", value=defaults["italic"], disabled=disabled_state),
        f"{label.lower()}_align": st.sidebar.selectbox(f"{label} Align", ["Left", "Center", "Right", "Justify"], index=defaults["align"], disabled=disabled_state),
    }

styles_config = {
    "title_size": st.sidebar.slider("Title Font Size", 14, 28, 22, disabled=disabled_state),
    "heading_size": st.sidebar.slider("Heading Font Size", 12, 20, 16, disabled=disabled_state),
    "subheading_size": st.sidebar.slider("Subheading Font Size", 10, 16, 12, disabled=disabled_state),
    "body_size": st.sidebar.slider("Body Font Size", 8, 14, 11, disabled=disabled_state),
}
styles_config.update(style_section("Title", {"bold": True, "italic": False, "align": 1}))
styles_config.update(style_section("Contact", {"bold": False, "italic": False, "align": 1}))
styles_config.update(style_section("Heading", {"bold": True, "italic": False, "align": 0}))
styles_config.update(style_section("Subheading", {"bold": True, "italic": False, "align": 0}))
styles_config.update(style_section("Body", {"bold": False, "italic": False, "align": 3}))

st.sidebar.checkbox("Live Preview", value=True, key="live_preview")

# Main layout: left = existing inputs/buttons, right = preview
left, right = st.columns([1, 1])

with left:
    # Inputs
    resume_template_text = st.text_area("Paste Your Resume Template", height=250, disabled=disabled_state)
    custom_prompt = st.text_area("Custom Prompt", height=200, value=DEFAULT_PROMPT, disabled=disabled_state)
    job_description = st.text_area("Paste the Job Description", height=300, disabled=disabled_state)
    include_tech = st.checkbox("Include Technologies per Position", value=True, disabled=disabled_state)
    output_format = st.selectbox("Output Format", ["DOCX", "PDF"], disabled=disabled_state)

    # Buttons on main page
    col1, col2, col3 = st.columns(3)

    with col1:
        st.button("✅ Generate Tailored Resume", on_click=generate_resume, disabled=disabled_state)

    with col2:
        if st.button("🔄 Reset App", key="reset_normal", disabled=st.session_state["loading"]):
            st.session_state.clear()
            st.rerun()

    with col3:
        # ✅ Only visible while loading; disappears automatically after completion
        if st.session_state.get("loading", False):
            if st.button("🛑 Stop Generating"):
                st.session_state["stop_requested"] = True
                st.session_state["loading"] = False
                st.toast("Generation stopped.", icon="🛑")
                st.rerun()

with right:
    st.subheader("🔎 Preview")
    if st.session_state.get("live_preview") and st.session_state.get("tailored_json"):
        preview_html = render_preview_html(st.session_state["tailored_json"], styles_config)
        st.markdown(preview_html, unsafe_allow_html=True)
    else:
        st.info("Generate once to enable live preview of your current styles.")


